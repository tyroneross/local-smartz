"""Report generation tool — create markdown, HTML, and DOCX reports."""

import json
from typing import Any
from pathlib import Path
from datetime import datetime

from langchain_core.tools import tool


@tool
def create_report(
    title: str,
    sections: list[dict] | str,
    output_path: str,
    format: str = "markdown",
    author: str | None = None,
    date: str | None = None,
    subtitle: str | None = None,
) -> str:
    """Generate a formatted report from structured sections.

    Args:
        title: Report title
        sections: List of sections with heading, content, and optional level.
                 Each section: {"heading": "...", "content": "...", "level": 2}
                 Can be list[dict] or JSON string
        output_path: Where to save the report
        format: Output format - markdown, html, or docx
        author: Optional report author
        date: Optional date (defaults to today)
        subtitle: Optional subtitle

    Returns:
        Success message with file path and stats
    """
    # Handle LangChain resilience - some models send JSON strings
    if isinstance(sections, str):
        try:
            sections = json.loads(sections)
        except json.JSONDecodeError:
            return f"Error: sections must be a valid JSON array, got: {sections[:100]}"

    if not isinstance(sections, list):
        return f"Error: sections must be a list, got {type(sections).__name__}"

    # Build metadata dict
    metadata = {}
    if author:
        metadata["author"] = author
    if date:
        metadata["date"] = date
    if subtitle:
        metadata["subtitle"] = subtitle

    # Ensure output directory exists
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    if format == "markdown":
        return _create_markdown_report(title, sections, output_path, metadata)
    elif format == "docx":
        return _create_docx_report(title, sections, output_path, metadata)
    elif format == "html":
        return _create_html_report(title, sections, output_path, metadata)
    else:
        return f"Error: Unsupported format: {format}. Use markdown, html, or docx"


@tool
def create_spreadsheet(
    data: list[dict] | str,
    output_path: str,
    sheet_name: str = "Sheet1",
) -> str:
    """Create an Excel spreadsheet from data.

    Args:
        data: List of dictionaries where keys are column headers.
              Can be list[dict] or JSON string.
              Example: [{"Name": "Alice", "Age": 30}, {"Name": "Bob", "Age": 25}]
        output_path: Where to save the .xlsx file
        sheet_name: Name for the worksheet (max 31 chars)

    Returns:
        Success message with file path and stats
    """
    # Handle LangChain resilience - some models send JSON strings
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError:
            return f"Error: data must be a valid JSON array, got: {data[:100]}"

    if not isinstance(data, list):
        return f"Error: data must be a list, got {type(data).__name__}"

    if not data:
        return "Error: data cannot be empty"

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    # Ensure output directory exists
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # Extract headers from first row
    first_row = data[0]
    if not isinstance(first_row, dict):
        return f"Error: each data item must be a dict, got {type(first_row).__name__}"

    headers = list(first_row.keys())

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]  # Excel sheet name limit

    # Write headers with bold font
    bold = Font(bold=True)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = bold

    # Write data rows
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, header in enumerate(headers, 1):
            value = row_data.get(header, "")
            ws.cell(row=row_idx, column=col_idx, value=value)

    # Auto-fit column widths
    for col_idx in range(1, len(headers) + 1):
        max_len = len(str(headers[col_idx - 1]))
        for row in ws.iter_rows(min_row=2, min_col=col_idx, max_col=col_idx):
            for cell in row:
                if cell.value is not None:
                    max_len = max(max_len, len(str(cell.value)))
        # Cap width at 50, add padding
        adjusted = min(max_len + 2, 50)
        ws.column_dimensions[get_column_letter(col_idx)].width = adjusted

    wb.save(output_path)

    # Register with artifacts system
    try:
        from localsmartz.artifacts import register
        register(
            path=output_path,
            format="xlsx",
            title=Path(output_path).stem,
            cwd=Path.cwd(),
        )
    except Exception:
        pass  # Artifact registration is best-effort

    saved = Path(output_path).resolve()
    if not saved.exists():
        return f"Error: Save completed but file not found at: {saved}"
    size = saved.stat().st_size
    if size == 0:
        return f"Error: File created but is empty: {saved}"

    return f"Spreadsheet saved: {saved} ({len(data)} rows, {len(headers)} columns, {size:,} bytes)"


def _register_artifact(output_path: str, fmt: str, title: str) -> None:
    """Best-effort artifact registration."""
    try:
        from localsmartz.artifacts import register
        register(path=output_path, format=fmt, title=title, cwd=Path.cwd())
    except Exception:
        pass


def _create_markdown_report(title: str, sections: list[dict], output_path: str, metadata: dict) -> str:
    lines: list[str] = []

    # Header
    lines.append(f"# {title}")
    if metadata.get("subtitle"):
        lines.append(f"### {metadata['subtitle']}")
    lines.append("")

    # Metadata
    meta_parts = []
    if metadata.get("author"):
        meta_parts.append(f"**Author**: {metadata['author']}")
    date_str = metadata.get("date", datetime.now().strftime("%Y-%m-%d"))
    meta_parts.append(f"**Date**: {date_str}")
    if meta_parts:
        lines.append(" | ".join(meta_parts))
        lines.append("")
        lines.append("---")
        lines.append("")

    # Sections
    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        level = section.get("level", 2)
        lines.append(f"{'#' * level} {heading}")
        lines.append("")
        lines.append(content)
        lines.append("")

    # Footer
    lines.append("---")
    lines.append(f"*Generated by Local Smartz on {datetime.now().strftime('%Y-%m-%d %H:%M')}*")

    text = "\n".join(lines)
    Path(output_path).write_text(text, encoding="utf-8")
    _register_artifact(output_path, "markdown", title)

    saved = Path(output_path).resolve()
    if not saved.exists():
        return f"Error: Save completed but file not found at: {saved}"
    size = saved.stat().st_size
    if size == 0:
        return f"Error: File created but is empty: {saved}"

    return f"Report saved: {saved} ({size:,} bytes, {len(sections)} sections)"


def _create_docx_report(title: str, sections: list[dict], output_path: str, metadata: dict) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_parts = []
    if metadata.get("author"):
        meta_parts.append(f"Author: {metadata['author']}")
    date_str = metadata.get("date", datetime.now().strftime("%Y-%m-%d"))
    meta_parts.append(f"Date: {date_str}")
    if metadata.get("subtitle"):
        meta_parts.append(metadata["subtitle"])
    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta_para.runs:
            run.font.size = Pt(10)
            run.font.italic = True

    doc.add_paragraph("_" * 60)  # Horizontal rule

    # Sections
    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        level = min(section.get("level", 2), 4)  # DOCX supports levels 0-4

        doc.add_heading(heading, level=level)

        # Process content: detect bullet lists vs paragraphs
        for para_text in content.split("\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            if stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif len(stripped) > 1 and stripped[0].isdigit() and stripped[1] in ".)":
                doc.add_paragraph(stripped[2:].strip(), style="List Number")
            else:
                doc.add_paragraph(stripped)

    # Footer
    doc.add_paragraph("_" * 60)
    footer = doc.add_paragraph(f"Generated by Local Smartz on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.italic = True

    doc.save(output_path)
    _register_artifact(output_path, "docx", title)

    saved = Path(output_path).resolve()
    if not saved.exists():
        return f"Error: Save completed but file not found at: {saved}"
    size = saved.stat().st_size
    if size == 0:
        return f"Error: File created but is empty: {saved}"

    return f"DOCX report saved: {saved} ({len(sections)} sections, {size:,} bytes)"


def _create_html_report(title: str, sections: list[dict], output_path: str, metadata: dict) -> str:
    date_str = metadata.get("date", datetime.now().strftime("%Y-%m-%d"))
    author = metadata.get("author", "")

    html_parts = [
        "<!DOCTYPE html>",
        "<html lang=\"en\">",
        "<head>",
        f"  <title>{_html_escape(title)}</title>",
        "  <meta charset=\"utf-8\">",
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">",
        "  <style>",
        "    body { font-family: -apple-system, system-ui, sans-serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; color: #1a1a1a; line-height: 1.6; }",
        "    h1 { border-bottom: 2px solid #333; padding-bottom: 0.5rem; }",
        "    h2 { margin-top: 2rem; color: #2c3e50; }",
        "    .meta { color: #666; font-size: 0.9rem; margin-bottom: 2rem; }",
        "    table { border-collapse: collapse; width: 100%; margin: 1rem 0; }",
        "    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
        "    th { background: #f5f5f5; }",
        "    blockquote { border-left: 3px solid #ccc; margin: 1rem 0; padding: 0.5rem 1rem; color: #555; }",
        "    code { background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }",
        "    pre { background: #f4f4f4; padding: 1rem; border-radius: 4px; overflow-x: auto; }",
        "    .footer { margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #ddd; color: #999; font-size: 0.8rem; }",
        "  </style>",
        "</head>",
        "<body>",
        f"  <h1>{_html_escape(title)}</h1>",
    ]

    if author or date_str:
        html_parts.append(f"  <p class=\"meta\">{_html_escape(author)}{' | ' if author else ''}{_html_escape(date_str)}</p>")

    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        level = section.get("level", 2)
        html_parts.append(f"  <h{level}>{_html_escape(heading)}</h{level}>")
        # Convert content paragraphs to HTML
        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                html_parts.append(f"  <p>{_html_escape(para)}</p>")

    html_parts.extend([
        f"  <div class=\"footer\">Generated by Local Smartz on {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>",
        "</body>",
        "</html>",
    ])

    html_text = "\n".join(html_parts)
    Path(output_path).write_text(html_text, encoding="utf-8")
    _register_artifact(output_path, "html", title)

    saved = Path(output_path).resolve()
    if not saved.exists():
        return f"Error: Save completed but file not found at: {saved}"
    size = saved.stat().st_size
    if size == 0:
        return f"Error: File created but is empty: {saved}"

    return f"HTML report saved: {saved} ({size:,} bytes)"


def _html_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
